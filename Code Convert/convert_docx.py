"""
convert_docx.py — Script Konversi File Word (.docx) ke Markdown (.md)
=====================================================================
Digunakan sebagai langkah pertama dalam pipeline pembuatan jurnal ilmiah.
Konversikan skripsi (DOCX) dan template jurnal (DOCX) ke format Markdown
agar dapat dibaca dan dianalisis oleh AI.

Cara pakai:
    1. Isi variabel INPUT_DOCX dan OUTPUT_MD di bagian bawah script
    2. Jalankan: python convert_docx.py

Dependensi:
    pip install python-docx
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn


def get_heading_level(paragraph):
    """Membaca level heading dari style paragraf."""
    style_name = paragraph.style.name
    if style_name.startswith('Heading'):
        try:
            level = int(style_name.split(' ')[-1])
            return level
        except:
            return 0
    return 0


def paragraph_to_md(paragraph):
    """Mengonversi satu paragraf Word ke string Markdown (bold/italic)."""
    text = ""

    heading_level = get_heading_level(paragraph)
    if heading_level > 0:
        full_text = paragraph.text.strip()
        if full_text:
            return '#' * heading_level + ' ' + full_text
        return ''

    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue
        if run.bold and run.italic:
            run_text = f'***{run_text}***'
        elif run.bold:
            run_text = f'**{run_text}**'
        elif run.italic:
            run_text = f'*{run_text}*'
        text += run_text

    return text.strip()


def table_to_md(table):
    """Mengonversi tabel Word ke format tabel Markdown."""
    lines = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        lines.append('| ' + ' | '.join(cells) + ' |')
        if i == 0:
            lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
    return '\n'.join(lines)


def docx_to_markdown(docx_path, md_path):
    """
    Fungsi utama konversi DOCX ke Markdown.
    
    Args:
        docx_path (str): Path ke file Word (.docx) input
        md_path   (str): Path ke file Markdown (.md) output
    """
    if not os.path.exists(docx_path):
        print(f"[ERROR] File tidak ditemukan: {docx_path}")
        return False

    doc = Document(docx_path)
    md_lines = []
    prev_was_list = False

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break

            if para is None:
                continue

            style_name = para.style.name.lower()
            text = para.text.strip()

            if not text:
                if md_lines and md_lines[-1] != '':
                    md_lines.append('')
                continue

            numPr = para._element.find(qn('w:pPr'))
            has_numbering = False
            ilvl = 0
            if numPr is not None:
                numEl = numPr.find(qn('w:numPr'))
                if numEl is not None:
                    has_numbering = True
                    ilvlEl = numEl.find(qn('w:ilvl'))
                    ilvl = int(ilvlEl.get(qn('w:val'), 0)) if ilvlEl is not None else 0

            if has_numbering:
                md_text = paragraph_to_md(para)
                if md_text:
                    indent = '  ' * ilvl
                    md_lines.append(f'{indent}- {md_text}')
                prev_was_list = True
            else:
                if prev_was_list:
                    md_lines.append('')
                prev_was_list = False

                md_text = paragraph_to_md(para)
                if md_text:
                    md_lines.append(md_text)

        elif tag == 'tbl':
            for tbl in doc.tables:
                if tbl._element is element:
                    if md_lines and md_lines[-1] != '':
                        md_lines.append('')
                    md_lines.append(table_to_md(tbl))
                    md_lines.append('')
                    break

    content = '\n'.join(md_lines)
    content = re.sub(r'\n{3,}', '\n\n', content)

    # Hapus konten Base64 yang sangat besar (opsional, aktifkan jika file terlalu besar)
    # content = re.sub(r'data:image/[^;]+;base64,[A-Za-z0-9+/=]{500,}', '[GAMBAR]', content)

    os.makedirs(os.path.dirname(md_path), exist_ok=True)

    with open(md_path, 'w', encoding='utf-8', errors='ignore') as f:
        f.write(content)

    print(f"[OK] Konversi berhasil!")
    print(f"     Input : {docx_path}")
    print(f"     Output: {md_path}")
    print(f"     Total baris: {len(md_lines)}")
    return True


# ============================================================
#  KONFIGURASI — Sesuaikan path di sini
# ============================================================
if __name__ == '__main__':

    # Konversi 1: Skripsi
    docx_to_markdown(
        docx_path = r'Skripsi\NAMA_SKRIPSI_ANDA.docx',       # <-- ganti nama file
        md_path   = r'Skripsi\skripsi.md'
    )

    # Konversi 2: Template Jurnal (jika dalam format DOCX)
    docx_to_markdown(
        docx_path = r'Template Jurnal\NAMA_TEMPLATE.docx',    # <-- ganti nama file
        md_path   = r'Template Jurnal\template.md'
    )
