"""
buat_jurnal.py — Script Generator Jurnal Format Word (.docx)
=============================================================
Script ini membangun file Word (.docx) berformat jurnal ilmiah dari kode Python
menggunakan library python-docx. Semua konten, tabel, dan gambar di-embed
langsung ke dalam file .docx sehingga portabel dan siap submit.

CARA PAKAI:
    1. Sesuaikan KONFIGURASI di bagian bawah script (judul, penulis, dll.)
    2. Isi konten jurnal pada masing-masing section (Section 1, 2, dst.)
    3. Jalankan: python buat_jurnal.py

DEPENDENSI:
    pip install python-docx

CATATAN:
    Spesifikasi format (margin, font, kolom) sudah diset untuk JTIIK.
    Sesuaikan bagian KONFIGURASI FORMAT jika menggunakan template jurnal lain.
"""

import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


# ==============================================================
#  KONFIGURASI JURNAL — Sesuaikan bagian ini
# ==============================================================

JUDUL_ID = "Judul Jurnal Dalam Bahasa Indonesia"
JUDUL_EN = "Journal Title in English"

PENULIS          = "Nama Lengkap Anda"
UNIVERSITAS      = "Nama Universitas"
PRODI_JURUSAN    = "Nama Program Studi / Jurusan"
EMAIL            = "email@anda.ac.id"

# Output file
OUTPUT_DOCX = r"Output\jurnal_output.docx"

# ==============================================================
#  KONFIGURASI FORMAT JURNAL (JTIIK — sesuaikan jika perlu)
# ==============================================================

PAGE_WIDTH_CM   = 21      # A4
PAGE_HEIGHT_CM  = 29.7    # A4
MARGIN_LEFT_CM  = 3
MARGIN_TOP_CM   = 3
MARGIN_RIGHT_CM = 2
MARGIN_BOT_CM   = 2
NUM_COLUMNS     = 2       # 2 = dua kolom; 1 = satu kolom

FONT_NORMAL     = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(10)
FONT_SIZE_TITLE  = Pt(12)
FONT_SIZE_TABLE  = Pt(8)
FONT_SIZE_CAPTION = Pt(8)


# ==============================================================
#  FUNGSI HELPER — Tidak perlu diubah
# ==============================================================

def setup_document():
    """Inisialisasi dokumen Word dengan format dasar."""
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = FONT_NORMAL
    style.font.size = FONT_SIZE_NORMAL

    section = doc.sections[0]
    section.page_width  = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin   = Cm(MARGIN_LEFT_CM)
    section.top_margin    = Cm(MARGIN_TOP_CM)
    section.right_margin  = Cm(MARGIN_RIGHT_CM)
    section.bottom_margin = Cm(MARGIN_BOT_CM)
    return doc


def add_2_column_section(doc):
    """Beralih ke layout 2 kolom untuk bagian body jurnal."""
    new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    sectPr = new_section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(NUM_COLUMNS))
    cols.set(qn('w:space'), '567')   # ~1cm
    sectPr.append(cols)


def add_text(doc, text, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             size=None, space_before=Pt(0), space_after=Pt(0)):
    """Menambahkan paragraf teks ke dokumen."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after  = space_after
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = size
    return p


def add_heading(doc, text):
    """Menambahkan heading bab (bold, uppercase)."""
    return add_text(doc, text, bold=True,
                    space_before=Pt(12), space_after=Pt(6))


def add_subheading(doc, text):
    """Menambahkan sub-heading (bold)."""
    return add_text(doc, text, bold=True,
                    space_before=Pt(6), space_after=Pt(6))


def add_caption(doc, text, is_table=False):
    """Menambahkan caption gambar atau tabel (8pt, center)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = FONT_SIZE_CAPTION
    run.bold = is_table   # Caption tabel biasanya di atas tabel, dibold
    return p


def add_image(doc, img_path, width_cm=7.5):
    """Menyisipkan gambar ke dokumen. Dilewati jika file tidak ditemukan."""
    if not os.path.exists(img_path):
        print(f"[SKIP] Gambar tidak ditemukan: {img_path}")
        add_text(doc, f"[Gambar: {os.path.basename(img_path)}]", italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))


def add_table(doc, headers, rows):
    """
    Membuat tabel dengan header dan data.

    Args:
        headers (list): List string untuk baris header
        rows    (list): List of lists untuk baris data
    """
    num_cols = len(headers)
    num_rows = len(rows) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Shading'

    # Header
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = FONT_SIZE_TABLE

    # Data rows
    for i, row_data in enumerate(rows, start=1):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = FONT_SIZE_TABLE

    return table


# ==============================================================
#  KONTEN JURNAL — Isi bagian ini sesuai isi jurnal Anda
# ==============================================================

def build_journal(doc):

    # ─── HEADER (1 kolom) ──────────────────────────────────────

    # Judul
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    r = p_title.add_run(JUDUL_ID)
    r.bold = True
    r.font.size = FONT_SIZE_TITLE

    add_text(doc, JUDUL_EN, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_text(doc, PENULIS, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, PRODI_JURUSAN + ", " + UNIVERSITAS,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Email: " + EMAIL,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # Abstrak Indonesia
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.add_run("ABSTRAK").bold = True
    p_abs.add_run("\n[Isi abstrak Bahasa Indonesia di sini. Sekitar 150–250 kata. "
                  "Uraikan: latar belakang masalah, tujuan, metode, hasil utama, dan kesimpulan.]")

    p_kw1 = doc.add_paragraph()
    p_kw1.paragraph_format.space_after = Pt(12)
    p_kw1.add_run("Kata kunci: ").bold = True
    p_kw1.add_run("kata kunci 1, kata kunci 2, kata kunci 3.")

    # Abstrak Inggris
    p_abs_en = doc.add_paragraph()
    p_abs_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs_en.add_run("ABSTRACT").bold = True
    r_en = p_abs_en.add_run("\n[English abstract here. Around 150–250 words. "
                             "Cover: background, objective, method, main results, and conclusion.]")
    r_en.italic = True

    p_kw2 = doc.add_paragraph()
    p_kw2.paragraph_format.space_after = Pt(24)
    p_kw2.add_run("Keywords: ").bold = True
    kw_run = p_kw2.add_run("keyword 1, keyword 2, keyword 3.")
    kw_run.italic = True

    # ─── PINDAH KE 2 KOLOM ────────────────────────────────────
    add_2_column_section(doc)

    # ─── 1. PENDAHULUAN ───────────────────────────────────────
    add_heading(doc, "1. PENDAHULUAN")
    add_text(doc, "[Paragraf 1: Uraikan konteks dan latar belakang masalah yang melatarbelakangi penelitian ini. "
                  "Jelaskan pentingnya masalah tersebut dan mengapa perlu diselesaikan.]")
    add_text(doc, "[Paragraf 2: Uraikan state-of-the-art dan penelitian terdahulu yang relevan. "
                  "Identifikasi gap penelitian yang ingin diisi oleh studi ini.]")
    add_text(doc, "[Paragraf 3: Nyatakan tujuan dan kontribusi penelitian secara eksplisit "
                  "(misal: 4 kontribusi utama secara poin bernomor).]")

    # ─── 2. TINJAUAN PUSTAKA ──────────────────────────────────
    add_heading(doc, "2. TINJAUAN PUSTAKA")
    add_text(doc, "[Uraikan 6–10 penelitian terkait dengan relevansinya terhadap penelitian ini. "
                  "Gunakan format: Nama et al. (tahun) melakukan/menemukan/mengusulkan ...]")

    # ─── 3. METODE PENELITIAN ─────────────────────────────────
    add_heading(doc, "3. METODE PENELITIAN")

    add_subheading(doc, "3.1 [Nama Sub-bagian Pertama]")
    add_text(doc, "[Deskripsikan dataset atau sumber data yang digunakan: jumlah data, "
                  "cara pengumpulan, anotasi, split train/val/test, augmentasi jika ada.]")

    add_image(doc, r"Hasil Model\[nama_gambar_arsitektur].png")
    add_caption(doc, "Gambar 1. [Keterangan gambar]")

    add_subheading(doc, "3.2 [Nama Sub-bagian Kedua]")
    add_text(doc, "[Deskripsikan arsitektur model, konfigurasi training, hyperparameter, "
                  "framework, dan hardware yang digunakan.]")

    add_subheading(doc, "3.3 Metrik Evaluasi")
    add_text(doc, "[Jelaskan metrik yang digunakan untuk mengukur performa sistem: "
                  "Precision, Recall, mAP, Accuracy, F1, dll.]")

    # ─── 4. HASIL DAN PEMBAHASAN ──────────────────────────────
    add_heading(doc, "4. HASIL DAN PEMBAHASAN")

    add_subheading(doc, "4.1 [Hasil Eksperimen Pertama]")
    add_text(doc, "[Uraikan hasil pengujian model pertama. Referensikan Tabel 1 dan Gambar 2.]")

    # Contoh tabel
    add_caption(doc, "Tabel 1. [Judul Tabel Perbandingan]", is_table=True)
    add_table(
        doc,
        headers=["Model", "Metrik 1", "Metrik 2", "Metrik 3"],
        rows=[
            ["Model A", "0.9X", "0.9X", "X ms"],
            ["Model B", "0.9X", "0.9X", "X ms"],
            ["Model C (Terpilih)", "0.9X", "0.9X", "X ms"],
        ]
    )
    doc.add_paragraph()  # spacing

    add_subheading(doc, "4.2 [Hasil Eksperimen Kedua]")
    add_text(doc, "[Uraikan hasil perbandingan kedua. Referensikan Tabel 2 dan Gambar 3.]")

    add_caption(doc, "Tabel 2. [Judul Tabel Perbandingan Kedua]", is_table=True)
    add_table(
        doc,
        headers=["Metode", "Akurasi (%)", "Metrik 2 (%)", "Keterangan"],
        rows=[
            ["Metode A", "9X.XX", "XX.XX", "Terbaik"],
            ["Metode B (Dipilih)", "6X.XX", "XX.XX", "Trade-off"],
            ["Metode C", "4X.XX", "XX.XX", "-"],
        ]
    )
    doc.add_paragraph()

    add_image(doc, r"Hasil Model\[nama_grafik_perbandingan].png")
    add_caption(doc, "Gambar 2. [Keterangan grafik perbandingan]")

    add_subheading(doc, "4.3 [Analisis Tambahan / Trade-off / Latensi]")
    add_text(doc, "[Uraikan analisis mendalam: trade-off keputusan desain, latensi sistem, "
                  "analisis error, atau pembahasan lainnya yang memperkuat kontribusi.]")

    # ─── 5. KESIMPULAN ────────────────────────────────────────
    add_heading(doc, "5. KESIMPULAN")
    add_text(doc, "[Ringkasan temuan utama: sebutkan nilai metrik terbaik, keputusan desain, "
                  "dan implikasi praktis dari hasil penelitian.]")
    add_text(doc, "[Saran pengembangan ke depan: sebutkan keterbatasan dan langkah berikutnya "
                  "yang bisa dilakukan oleh peneliti lain.]")

    # ─── DAFTAR PUSTAKA ───────────────────────────────────────
    add_heading(doc, "DAFTAR PUSTAKA")
    refs = [
        "Penulis, A., & Penulis, B. (Tahun). Judul artikel. Nama Jurnal, Volume(Issue), Halaman.",
        "Penulis, C. (Tahun). Judul buku. Penerbit.",
        "Penulis, D., et al. (Tahun). Judul conference paper. Nama Konferensi, Halaman.",
        # Tambahkan referensi lainnya di sini
    ]
    for ref in refs:
        add_text(doc, ref, space_after=Pt(6))


# ==============================================================
#  MAIN
# ==============================================================
if __name__ == '__main__':
    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)

    doc = setup_document()
    build_journal(doc)
    doc.save(OUTPUT_DOCX)

    print(f"[OK] File jurnal berhasil dibuat: {OUTPUT_DOCX}")
